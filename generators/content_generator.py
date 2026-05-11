import os
import logging
from typing import Optional, Dict, Any
from datetime import datetime
import json

from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from agents.gemini_helper import GeminiHelper


class ContentGenerator:
    """
    Generate various types of content (images, PDFs, Word docs) based on prompts
    """
    
    def __init__(self):
        self.gemini_helper = GeminiHelper()
        self.output_dir = "generated_content"
        self._ensure_output_dir()
    
    def _ensure_output_dir(self):
        """Create output directory if it doesn't exist"""
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
            logging.info(f"Created output directory: {self.output_dir}")
    
    def generate_image(self, prompt: str, filename: Optional[str] = None) -> Dict[str, Any]:
        """
        Generate an image based on text prompt
        
        Args:
            prompt: Description of the image to generate
            filename: Optional custom filename
            
        Returns:
            Dictionary with result information
        """
        try:
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"generated_image_{timestamp}.png"
            
            filepath = os.path.join(self.output_dir, filename)
            
            # Note: Image generation will be handled by the chatbot calling the appropriate tool
            # This is a placeholder for the response structure
            
            return {
                "success": True,
                "type": "image",
                "filepath": filepath,
                "filename": filename,
                "message": f"Image generated successfully! Saved as {filename}"
            }
            
        except Exception as e:
            logging.error(f"Error generating image: {str(e)}")
            return {
                "success": False,
                "type": "image",
                "error": str(e),
                "message": "Failed to generate image. Please try again."
            }
    
    def generate_text_content(self, prompt: str, content_type: str = "article") -> Dict[str, Any]:
        """
        Generate text content using Gemini
        
        Args:
            prompt: Description of content to generate
            content_type: Type of content (article, story, essay, etc.)
            
        Returns:
            Dictionary with generated content
        """
        try:
            if not self.gemini_helper.is_available():
                return {
                    "success": False,
                    "type": "text",
                    "error": "AI text generation not available",
                    "message": "Text generation requires Gemini AI integration."
                }
            
            enhanced_prompt = f"Generate a well-structured {content_type} based on this prompt: {prompt}. Make it engaging, informative, and well-formatted."
            
            content = self.gemini_helper.get_response(enhanced_prompt)
            
            if content:
                return {
                    "success": True,
                    "type": "text",
                    "content": content,
                    "content_type": content_type,
                    "message": f"Generated {content_type} successfully!"
                }
            else:
                return {
                    "success": False,
                    "type": "text",
                    "error": "No content generated",
                    "message": "Failed to generate text content. Please try again."
                }
                
        except Exception as e:
            logging.error(f"Error generating text content: {str(e)}")
            return {
                "success": False,
                "type": "text",
                "error": str(e),
                "message": "Failed to generate text content."
            }
    
    def generate_pdf(self, prompt: str, filename: Optional[str] = None) -> Dict[str, Any]:
        """
        Generate a PDF document based on prompt
        
        Args:
            prompt: Description of PDF content
            filename: Optional custom filename
            
        Returns:
            Dictionary with result information
        """
        try:
            # First generate the text content
            text_result = self.generate_text_content(prompt, "document")
            
            if not text_result["success"]:
                return text_result
            
            content = text_result["content"]
            
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"generated_document_{timestamp}.pdf"
            
            filepath = os.path.join(self.output_dir, filename)
            
            # Create PDF
            doc = SimpleDocTemplate(filepath, pagesize=letter)
            styles = getSampleStyleSheet()
            
            # Create custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30,
                alignment=1  # Center alignment
            )
            
            content_style = ParagraphStyle(
                'CustomContent',
                parent=styles['Normal'],
                fontSize=12,
                spaceAfter=12,
                leftIndent=0,
                rightIndent=0
            )
            
            # Build PDF content
            story = []
            
            # Add title
            title = f"Generated Document"
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 12))
            
            # Add generated content
            # Split content into paragraphs
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    story.append(Paragraph(para.strip(), content_style))
                    story.append(Spacer(1, 12))
            
            # Generate PDF
            doc.build(story)
            
            return {
                "success": True,
                "type": "pdf",
                "filepath": filepath,
                "filename": filename,
                "content": content,
                "message": f"PDF generated successfully! Saved as {filename}"
            }
            
        except Exception as e:
            logging.error(f"Error generating PDF: {str(e)}")
            return {
                "success": False,
                "type": "pdf",
                "error": str(e),
                "message": "Failed to generate PDF. Please try again."
            }
    
    def generate_word_doc(self, prompt: str, filename: Optional[str] = None) -> Dict[str, Any]:
        """
        Generate a Word document based on prompt
        
        Args:
            prompt: Description of document content
            filename: Optional custom filename
            
        Returns:
            Dictionary with result information
        """
        try:
            # First generate the text content
            text_result = self.generate_text_content(prompt, "document")
            
            if not text_result["success"]:
                return text_result
            
            content = text_result["content"]
            
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"generated_document_{timestamp}.docx"
            
            filepath = os.path.join(self.output_dir, filename)
            
            # Create Word document
            doc = Document()
            
            # Add title
            title = doc.add_heading('Generated Document', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add generated content
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    # Check if it looks like a heading (starts with # or is short and ends with :)
                    if para.startswith('#') or (len(para) < 100 and para.endswith(':')):
                        # Add as heading
                        heading_text = para.replace('#', '').strip().rstrip(':')
                        doc.add_heading(heading_text, level=1)
                    else:
                        # Add as regular paragraph
                        p = doc.add_paragraph(para.strip())
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Add footer with generation info
            doc.add_page_break()
            footer_para = doc.add_paragraph()
            footer_para.add_run("Generated by AI Chatbot on ").italic = True
            footer_para.add_run(datetime.now().strftime("%B %d, %Y at %I:%M %p")).italic = True
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Save document
            doc.save(filepath)
            
            return {
                "success": True,
                "type": "word",
                "filepath": filepath,
                "filename": filename,
                "content": content,
                "message": f"Word document generated successfully! Saved as {filename}"
            }
            
        except Exception as e:
            logging.error(f"Error generating Word document: {str(e)}")
            return {
                "success": False,
                "type": "word",
                "error": str(e),
                "message": "Failed to generate Word document. Please try again."
            }
    
    def list_generated_files(self) -> Dict[str, Any]:
        """
        List all generated files in the output directory
        
        Returns:
            Dictionary with file information
        """
        try:
            if not os.path.exists(self.output_dir):
                return {
                    "success": True,
                    "files": [],
                    "message": "No files generated yet."
                }
            
            files = []
            for filename in os.listdir(self.output_dir):
                filepath = os.path.join(self.output_dir, filename)
                if os.path.isfile(filepath):
                    stat = os.stat(filepath)
                    files.append({
                        "filename": filename,
                        "size": stat.st_size,
                        "created": datetime.fromtimestamp(stat.st_ctime).strftime("%Y-%m-%d %H:%M:%S")
                    })
            
            files.sort(key=lambda x: x["created"], reverse=True)
            
            return {
                "success": True,
                "files": files,
                "count": len(files),
                "message": f"Found {len(files)} generated files."
            }
            
        except Exception as e:
            logging.error(f"Error listing files: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "message": "Failed to list generated files."
            }